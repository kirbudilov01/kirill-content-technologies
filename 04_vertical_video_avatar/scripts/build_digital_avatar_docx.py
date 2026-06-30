from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "digital-avatar-ready-systems-report.md"
DOCX_PATH = ROOT / "digital-avatar-ready-systems-report.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, url, text=None):
    text = text or url
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_rich_text(paragraph, text):
    pattern = re.compile(r"\[([^\]]+)\]\((https?://[^)]+)\)|(https?://\S+)|(`[^`]+`)|(\*\*[^*]+\*\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        if match.group(1) and match.group(2):
            add_hyperlink(paragraph, match.group(2), match.group(1))
        elif match.group(3):
            add_hyperlink(paragraph, match.group(3), match.group(3))
        elif match.group(4):
            run = paragraph.add_run(match.group(4)[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        elif match.group(5):
            run = paragraph.add_run(match.group(5)[2:-2])
            run.bold = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def is_table_start(lines, i):
    return (
        i + 1 < len(lines)
        and lines[i].strip().startswith("|")
        and lines[i + 1].strip().startswith("|")
        and set(lines[i + 1].strip().replace("|", "").replace(":", "").replace("-", "").strip()) == set()
    )


def split_table_row(line):
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph()
    title_run = title.add_run("Digital Avatar Systems For AtlasRepo Content Factory")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(22)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string("0B2545")
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.add_paragraph()
    subtitle.add_run("RunPod-ready research brief for partner review").italic = True
    subtitle.paragraph_format.space_after = Pt(12)

    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_buf = []
    skip_first_h1 = True

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                p = doc.add_paragraph()
                p.style = styles["Normal"]
                for code_line in code_buf:
                    run = p.add_run(code_line + ("\n" if code_line != code_buf[-1] else ""))
                    run.font.name = "Consolas"
                    run.font.size = Pt(9.5)
                p.paragraph_format.left_indent = Inches(0.25)
                in_code = False
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if skip_first_h1 and line.startswith("# "):
            skip_first_h1 = False
            i += 1
            continue

        if is_table_start(lines, i):
            header = split_table_row(lines[i])
            rows = []
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1

            table = doc.add_table(rows=1, cols=len(header))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            table.autofit = True
            set_cell_margins(table)

            for idx, text in enumerate(header):
                cell = table.rows[0].cells[idx]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_shading(cell, "E8EEF5")
                p = cell.paragraphs[0]
                run = p.add_run(text)
                run.bold = True

            for row in rows:
                cells = table.add_row().cells
                for idx, text in enumerate(row):
                    cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    add_rich_text(cells[idx].paragraphs[0], text)
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, line[2:])
        elif re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_rich_text(p, re.sub(r"^\d+\.\s+", "", line))
        else:
            p = doc.add_paragraph()
            add_rich_text(p, line)
        i += 1

    footer = section.footer.paragraphs[0]
    footer.text = "AtlasRepo avatar systems research | 2026-06-27"
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor.from_string("666666")

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
