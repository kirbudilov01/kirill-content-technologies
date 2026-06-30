from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt, RGBColor


OUT = Path("research/youtube/outputs/first_batch")
OUT.mkdir(parents=True, exist_ok=True)


def setup_doc(title: str, subtitle: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    heading_specs = {
        "Heading 1": (20, RGBColor(0, 0, 0), 20, 6),
        "Heading 2": (16, RGBColor(0, 0, 0), 18, 6),
        "Heading 3": (14, RGBColor(67, 67, 67), 16, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        st = styles[name]
        st.font.name = "Arial"
        st.font.size = Pt(size)
        st.font.bold = False
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.15

    for style_name, size, after in [
        ("ScriptBody", 11.5, 10),
        ("BriefBody", 10.8, 7),
        ("SmallNote", 10, 5),
    ]:
        if style_name not in styles:
            s = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            s.base_style = styles["Normal"]
        s = styles[style_name]
        s.font.name = "Arial"
        s.font.size = Pt(size)
        s.paragraph_format.line_spacing = 1.15
        s.paragraph_format.space_after = Pt(after)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.font.name = "Arial"
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = False

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(subtitle)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = RGBColor(85, 85, 85)
    return doc


def add_script(doc: Document, title: str, meta: str, text: str):
    doc.add_heading(title, level=1)
    p = doc.add_paragraph(meta, style="SmallNote")
    p.runs[0].italic = True
    for para in [x.strip() for x in text.strip().split("\n\n") if x.strip()]:
        doc.add_paragraph(para, style="ScriptBody")


def add_brief_section(doc: Document, title: str, lines):
    doc.add_heading(title, level=2)
    for line in lines:
        doc.add_paragraph(line, style="BriefBody")


video1 = """
Imagine trying to build a one-person AI company in 2026. You do not have a research team, an editor, a developer, a CRM person, a social media manager, or someone who clicks through boring websites all day. But you do have GitHub. And that changes the entire game.

In this video, I am not going to give you another random list of “cool AI tools.” That is not useful. I want to show you ten open-source repositories that each replace one real job inside a small company. One repo can use websites for you. One can help create short videos. One can run AI models locally. One gives you your own AI workspace. One helps you build AI apps. One creates visual AI workflows. One works inside codebases. One turns agents into teammates. One gives you an open-source CRM. And one takes the content you make and pushes it out into the world.

The first repo is browser-use. Here is the human version of the problem: imagine you need to check forty competitor websites, fill out forms, compare pricing pages, or test the same user flow over and over again. Most people still think AI is just a chat box that writes text. Browser-use is different. It lets AI operate websites. It can open a browser, click, type, navigate, and return a result. That is why I would put it first in the video. Not because it is the whole company, but because the moment you see an AI actually using a website, the idea clicks instantly.

If I were showing this on screen, I would not start with the README. I would start with the task. “Go to this website, find the relevant page, extract the useful information.” Then I would show the agent moving through the browser. That is the save-worthy moment. The audience does not need a lecture about automation frameworks. They need to see a simple job that used to take a human, now being handled by an AI worker.

The second repo is MoneyPrinterTurbo. The name sounds ridiculous, but the use case is easy to understand. Imagine you have an idea for a short video, but no editor, no motion designer, and no time to manually assemble every version. This repo points toward AI-generated Shorts, Reels, and TikToks. I would be careful with the promise here. I would not say it magically replaces a great editor. I would say this: if you are building a content engine, you need to understand where open-source video generation is going, because the first draft of a short video is starting to become software.

On screen, the only thing that matters is the before and after. Before: a topic. After: a generated video draft, or at least a visible pipeline that moves from input to output. If the output is rough, say that. That actually builds trust. The point is not “this is perfect.” The point is “this is already real enough that you should be watching it.”

The third repo is Ollama. Here is the simple context: every AI workflow becomes fragile if it only works by renting intelligence from someone else’s cloud. Ollama gives you a local AI engine. You can run models on your own machine or your own infrastructure, and then connect other tools to that foundation. For a non-technical viewer, I would explain it like this: Ollama is not the whole app. It is the engine under the hood.

I would show one quick terminal moment: install, pull a model, run a prompt, get a response. No benchmark rabbit hole. No deep model comparison. The retention move is simple: “Look, the model is running here, not inside somebody else’s website.” That is enough.

The fourth repo is Open WebUI. If Ollama is the engine, Open WebUI is the cockpit. Now imagine you can run local or cloud models, but you also want a normal workspace that feels closer to ChatGPT: chats, model selection, settings, files, and a clean interface. That is what makes Open WebUI powerful for a broader audience. It turns local AI from a terminal trick into something a creator, founder, or small team can actually use every day.

This is the first natural moment to mention AtlasRepo. We found these projects while reviewing thousands of open-source tools in AtlasRepo, our free scanner for open-source projects and AI solutions. The point of AtlasRepo is not to create a giant graveyard of links. The point is to turn GitHub chaos into a map: AI tools, agents, automation, content systems, devtools, self-hosted alternatives, and open-source startups. So when I show a repo here, I am not saying “I saw one tweet and copied it.” I am saying this belongs to a role inside a real workflow. If you want to scan this kind of stuff yourself, the free AtlasRepo link is in the channel header.

The fifth repo is Dify. Here is the context: imagine you have an idea for an AI product, but you do not want to start from an empty folder and rebuild chat, RAG, workflows, knowledge bases, tool calls, and deployment from scratch. Dify gives you a platform for building AI apps faster. This is where the video should shift from “AI tools are cool” to “I can actually build something with this.”

The best demo is not a menu tour. The best demo is a tiny app. Create an assistant, connect a knowledge source, publish a usable screen, and show the result. The audience should feel: “Wait, I could build a support bot, a research assistant, or an internal knowledge tool this weekend.”

The sixth repo is ComfyUI. Imagine needing thumbnails, product mockups, ad creatives, visual experiments, and assets for content every single week. ComfyUI is not just an image generator. It is a visual workflow system for AI images and media. The easy explanation is: it is like a workflow board for AI visuals, where prompts, models, and outputs are connected through nodes.

This one is in the list because it is instantly visual. A node graph looks interesting. The output looks interesting. And creators understand the pain immediately: good visuals are expensive, slow, and annoying to produce consistently. ComfyUI is one of those repos people save because even if they do not use it today, they can feel the leverage.

The seventh repo is OpenHands. Here is the human version: your repo has ten annoying tasks, five small bugs, and one feature you keep postponing because it requires looking through files and making careful changes. OpenHands is an AI software engineer that can work inside repositories. I would not frame this as “fire your developers.” That sounds fake. I would frame it as: stop writing every mechanical line yourself.

For the screen recording, show the task, the files, the changes, and the result. The key is to make it concrete. “Change this part of the app.” “Add this small feature.” “Fix this issue.” If the viewer only sees a dashboard, they will not care. If they see an AI worker touch a real repo and produce a real diff, they will lean in.

The eighth repo is Multica. This is where the list gets more interesting than another generic AI-agent roundup. Imagine one AI helper is useful, but not enough. Now imagine assigning tasks to AI teammates, watching progress, and managing coding agents more like a small team. That is the promise of Multica. It fits the one-person company idea perfectly, because the future solo founder is not really alone. They are managing a stack of AI workers.

This is a “test, but pay attention” repo. I would not overpromise it. I would say: this is the kind of project that shows where the market is going. The interface matters a lot here. The partner needs to capture the dashboard, tasks, agents, progress, and any visible proof that this is more than a chat window.

The ninth repo is Twenty CRM. And this is important because a real company is not just AI demos. A real company has leads, customers, deals, follow-ups, and a pipeline. Imagine your business contacts are scattered across DMs, spreadsheets, notes, email, and random Notion pages. Twenty is an open-source CRM, a modern alternative to Salesforce-style software. It makes the list feel like an actual operating stack, not a bag of AI toys.

On screen, this should be very simple: contacts, companies, pipeline, deals, dashboard. The hook is not “CRM is exciting.” The hook is “your AI company still needs a place where money and customers are tracked.”

The tenth repo is Postiz. And this is the one I want to explain with the most normal-person context possible. Imagine that in 2026 it is no longer enough to record one video and hope the algorithm magically finds it. You need to publish versions, schedule posts, distribute clips, and show up on multiple platforms. Now imagine one open-source repo, with a free/self-hostable path, can become the distribution layer for that content. Meet Postiz.

Postiz is an open-source social media scheduling tool. Calendar, posts, queues, channels, distribution. It is not here because scheduling is sexy. It is here because content without distribution is invisible. In a one-person company, Postiz is the part of the system that turns “I made something” into “people might actually see it.”

Around seventy percent into the video, I would bring AtlasRepo back naturally. The big idea is not that these ten repos are the final list forever. The big idea is the method. If you want to build a product, a channel, an agent workflow, or a one-person company, you need a radar for open-source leverage. AtlasRepo is where we are building that radar: finding projects, sorting them by role, and turning them into practical stacks you can actually test. Again, the free scanner is linked from the channel header if you want to look for your own stack.

If I had to turn this into a weekend plan, I would not install all ten at once. I would pick one lane. If you want the fastest wow, try browser-use or ComfyUI. If you want your own AI base, set up Ollama and Open WebUI. If you want to build a product, open Dify. If you want the future of coding work, watch OpenHands and Multica. If you want the business side to stop being chaos, look at Twenty. And if you are already making content, Postiz is the repo that makes the whole thing feel less like random posting and more like a system.

That is the real takeaway. GitHub is starting to look less like a code archive and more like a hiring page for tiny teams. You are not just saving links. You are collecting roles. Browser worker. Video generator. Local AI engine. Workspace. App builder. Visual pipeline. Coding worker. Agent team. CRM. Distribution. Put the right roles together, and a one-person company stops sounding like a motivational quote and starts looking like an operating system.
"""


video2 = """
In this video, I am not just reviewing tools. I am going to build a tiny micro-SaaS idea from open-source projects in one evening. And the product is something we actually need: an Open-Source Tool Radar for Content Creators.

Here is the problem. Imagine waking up and needing to create content for a tech channel, a startup, or a personal brand. Most people start with a blank document and ask, “What should I post today?” That is a terrible place to start. The internet is already giving you signals every day: new GitHub repos, new agents, new AI video tools, new automation stacks, new open-source alternatives. The opportunity is not inventing ideas from nothing. The opportunity is turning those signals into a repeatable content system.

So the micro-SaaS is simple. You enter a niche, like “AI video tools,” “content automation,” or “solo founder stack.” The product returns ten GitHub repositories grouped by role. For each repo, it gives you a plain-English explanation, a save/test/skip verdict, three short-form content angles, one long-form video idea, and a short screencast brief for a partner.

That last part matters. A normal directory gives you links. A useful creator tool gives you production-ready output. I do not want another dashboard that looks impressive and does nothing. I want a system where the output is: “Here is the repo, here is why it matters, here is what to record, here is the hook, here is the script direction, now go shoot.”

The first screen is the niche input. I type “content factory.” The system should not return generic AI tools. It should understand roles. Research. Video generation. Editing. Publishing. Knowledge base. Automation. Distribution. This is the difference between a list and a stack. A list says “here are ten tools.” A stack says “this one finds ideas, this one creates videos, this one schedules content, this one stores research, this one automates the boring parts.”

For the first demo version, I would not overbuild it. One evening means we keep the MVP honest. Use AtlasRepo as the source of curated open-source projects. Start with a structured dataset: repo name, GitHub URL, category, role, short description, why creators care, and demo availability. Then build a simple interface that filters by niche and generates content cards.

This is the first natural AtlasRepo mention. The reason AtlasRepo matters here is that the micro-SaaS needs good input. If the source data is weak, the content plan is weak. AtlasRepo gives us a way to discover and organize open-source projects by use case: AI agents, automation, content tools, devtools, self-hosted SaaS, and open-source startups. It is also the free scanner I would point people to from the channel header, because this whole workflow starts with finding better raw material. The micro-SaaS is basically an action layer on top of that discovery engine.

Now let’s define the actual product. The name can be “Repo Radar” or “Creator Stack Finder.” The promise is: find the open-source tools worth making content about before everyone else turns them into recycled Twitter threads. The user enters a topic. The product returns a creator-ready brief.

A good result card should have seven fields. First: repository name and GitHub URL. Second: what it replaces in normal human language. Third: why the viewer should care. Fourth: visual proof to capture on screen. Fifth: risk or limitation. Sixth: short-form hooks. Seventh: long-form angle.

Let’s say the repo is Postiz. A weak card says: “Open-source social media scheduler.” A strong card says: “Imagine recording one video, then using one open-source repo to schedule and distribute the content across platforms instead of manually posting everywhere. That is the context. Show calendar, composer, queue, and connected channels.” That is production-ready. The partner can record it. I can read it. The viewer understands it in two seconds.

Let’s say the repo is browser-use. A weak card says: “AI browser automation.” A strong card says: “Imagine you need to check forty competitor websites and fill the same boring form again and again. This repo lets AI operate a browser like a human worker.” Now we have a hook, a demo, and a reason to care.

This is where the micro-SaaS becomes more than a directory. It becomes a content operating system. Partner opens the brief and records GitHub, README, demo, interface, limitation, and result. I open the script and record voiceover. Editor gets a clean sequence: problem, repo, proof, verdict. Postiz can distribute the final clip. AtlasRepo keeps feeding the system with new projects.

The second screen should be the repo stack. For “content factory,” I want to see a mix like Postiz for distribution, MoneyPrinterTurbo for AI short-video generation, OpenMontage for deeper video production, browser-use for research automation, Dify for building internal AI assistants, Open WebUI for AI workspace, Ollama for local models, ComfyUI for visuals, Twenty for CRM, and maybe Papra or Memos for research storage. The exact list can change. The important thing is that each project has a role.

The third screen is the content output. For each repo, generate three Shorts. One pain hook, one wow hook, one comparison hook. For Postiz: “You do not need to manually post the same content everywhere anymore.” For MoneyPrinterTurbo: “What if the first draft of your next Short came from a GitHub repo?” For browser-use: “This AI does not just answer. It clicks.”

Around seventy percent into the video, I would make the second AtlasRepo mention. This is where the product vision expands. Today, we built a tiny internal tool for our own content workflow. But the bigger idea is that AtlasRepo can become the place where open-source discovery turns into action. Not just “save this repo,” but “build with it, make content about it, compare it, test it, and publish the result.” If the viewer wants the starting point, the free open-source scanner is linked in the channel header.

The final product does not need to be perfect. It needs to make one moment feel obvious: a niche goes in, a usable production plan comes out. That is the before and after. Before, I am staring at a blank page. After, I have repos, angles, scripts, and recording tasks.

And that is why I like this micro-SaaS idea. It is not trying to be another beautiful AI dashboard that nobody opens twice. It has a job. It turns open-source discovery into a shooting schedule. If we can make that work for our own channel first, then it becomes much easier to imagine it as a public AtlasRepo feature later.
"""


video3 = """
In this video, I want to show you a different way to find open-source tools. Not by searching “best AI tools” and getting the same recycled list. Not by scrolling Twitter for whatever is loud this week. I want to start with a job to be done: build a content factory. Then I want to use AtlasRepo to find three GitHub repositories that actually fit that job.

The job is simple. I need to research ideas, create or assemble short videos, and distribute the final content across platforms. If a repo does not help with one of those steps, it might be interesting, but it is not gold for this specific stack. That is the filter.

The first gold repo is Postiz. And I want to introduce it with the right context. Imagine it is no longer enough to record one video and upload it once. You need clips, variations, scheduled posts, and distribution across platforms. Now imagine one open-source repo, with a free/self-hostable path, can become the layer that helps publish and schedule that content. That is why Postiz is more than a social media calendar. It is the distribution layer of a content factory.

On screen, I would show the GitHub repo first, then the product interface: calendar, composer, queue, and connected channels. The viewer should understand immediately: this is where finished content goes after it is created. The save reason is obvious. If you are a creator, founder, agency, or small team, distribution is not optional.

The second gold repo is MoneyPrinterTurbo. The context here is not “get rich with AI videos.” That sounds cheap. The stronger frame is: imagine you have ten content ideas, but you do not have the time or editing capacity to turn each one into a short-video draft. MoneyPrinterTurbo points toward a world where the first version of a Short, Reel, or TikTok can be generated by software.

The important thing is to show it honestly. If the result is rough, say it is rough. That does not kill the video. It makes the video credible. The reason it is gold is not that it produces perfect cinema. The reason it is gold is that AI video generation is becoming open-source, scriptable, and repeatable.

This is the first AtlasRepo moment. AtlasRepo is useful here because we are not just collecting links. We are asking: what role does this repo play in a real workflow? For this video, the workflow is content factory. So a project is only “gold” if it helps with research, production, editing, storage, automation, or distribution. If you want to run the same kind of scan yourself, the free AtlasRepo scanner is linked from the channel header.

The third gold repo is OpenMontage. This one is a little more advanced, and that is why I would frame it carefully. Imagine a video production pipeline where agents help plan, assemble, and structure video assets instead of everything living inside manual editing chaos. OpenMontage points in that direction. It may not be as instantly simple as Postiz, but it is the kind of repo that makes you stop and ask, “Wait, is open-source video production about to get much more serious?”

For the partner brief, this one needs strong visual proof. Show the GitHub repo, docs, workflow, demo assets, examples, or anything that communicates the production pipeline. If the demo is too hard to run, use README visuals and explain the concept clearly. The viewer does not need every implementation detail. They need to understand why it matters.

What makes these three repos work together is that they do not duplicate each other. Postiz distributes. MoneyPrinterTurbo generates short-video drafts. OpenMontage points toward deeper video production workflows. That gives the video a clean structure: publishing layer, short-form generation layer, advanced production layer.

Around seventy percent into the video, I would bring AtlasRepo back with the repeatable method. If you want to find gold repos, stop searching broad hype terms. Search by role. “I need to publish content.” “I need to generate short videos.” “I need to automate research.” “I need an open-source CRM.” “I need a local AI stack.” Good discovery starts with a job, not a buzzword. That is exactly why we keep pointing people to the free scanner in the channel header.

My read is simple. Postiz is the practical one. MoneyPrinterTurbo is the curiosity click. OpenMontage is the future-facing bet. I would not use the same rating scale for all three, because they do different jobs. One helps you publish. One helps you generate drafts. One hints at where video production may go next.

That is the format I want to repeat. Not “here are three random tools,” but “here is one job, and here are the open-source projects that can do parts of it.” Next time it could be a sales stack, a local AI stack, a solo founder stack, or an agent stack. AtlasRepo becomes the map, the channel becomes the field test, and the viewer gets the part that matters: what is actually worth opening after the video ends.
"""


shorts = [
    ("browser-use", "Did you know you can stop doing repetitive website work by hand? I recently found a GitHub repo called browser-use. It lets AI open a browser, click, type, navigate pages, and bring back results from a real web task. So instead of manually checking competitors, filling forms, or testing the same flow again, you can start turning that into an AI worker. Is this actually useful, or still too early? Tell me in the comments. And if you want more open-source solutions like this, check AtlasRepo in the channel header."),
    ("MoneyPrinterTurbo", "Did you know your first Short draft can already come from an open-source repo? I recently found MoneyPrinterTurbo. You give it a topic, and it tries to generate a short-video draft for Shorts, Reels, or TikToks. I am not saying it replaces a great editor, but this is the kind of category creators should watch before it becomes normal. Would you use something like this for content? Drop your take in the comments. More tools like this are on AtlasRepo in the channel header."),
    ("Ollama", "Did you know you do not have to run every AI workflow through paid cloud APIs? I recently found Ollama, one of the most important open-source repos for running AI models locally. You pull a model, run it on your machine, and use it as the engine for other AI tools. Not the flashiest demo, but a serious foundation. Would you rather use local AI or cloud AI? Write it in the comments. And if you want more open-source AI solutions, AtlasRepo is in the channel header."),
    ("Open WebUI", "Did you know local AI does not have to feel like a terminal-only hacker project? I recently found Open WebUI. It gives local or cloud models a clean ChatGPT-style workspace: chats, model selection, settings, and a real interface. Basically, Ollama gives you the engine, Open WebUI gives you the room where you actually work. Would you use your own AI workspace instead of another paid AI app? Tell me in the comments. More open-source solutions are on AtlasRepo in the channel header."),
    ("Dify", "Did you know you can build an AI app without starting from an empty folder? I recently found Dify, an open-source repo for assistants, workflows, knowledge bases, RAG, and app screens. The big thing here is speed: you can go from idea to usable prototype without rebuilding every boring layer yourself. Would you use this for an internal tool or a real SaaS? Let me know in the comments. And if you want more repos like this, AtlasRepo is linked in the channel header."),
    ("ComfyUI", "Did you know some open-source tools are basically production studios now? I recently found ComfyUI, and it looks like a control room for AI visuals. You connect prompts, models, nodes, and outputs into repeatable image workflows. For thumbnails, product shots, ad concepts, or visual experiments, this is way more than a simple image generator. Is this too complex, or exactly where AI design is going? Tell me in the comments. More creative open-source tools are on AtlasRepo in the channel header."),
    ("OpenHands", "Did you know an open-source repo can already act like a junior AI developer inside your codebase? I recently found OpenHands. It can inspect a project, make changes, and show you what happened. You still review the diff, obviously, but if it handles boring bugs, cleanup, and small features, that is real leverage. Would you trust an AI engineer with a small task in your repo? Comment yes or no. I collect more tools like this on AtlasRepo, linked in the channel header."),
    ("Multica", "Did you know the next step after one AI assistant might be a whole team of AI agents? I recently found Multica, a repo that treats coding agents more like teammates: assign tasks, watch progress, coordinate work. This is why the one-person company idea is getting interesting. You are not pretending to have a huge team; you are learning to manage AI workers. Is this exciting or terrifying? Tell me in the comments. More agent tools are on AtlasRepo in the channel header."),
    ("Twenty CRM", "Did you know one of the most useful open-source repos for an AI business is not even an AI tool? I recently found Twenty, an open-source CRM for contacts, companies, deals, and pipeline. If your leads live in DMs, spreadsheets, notes, and memory, you do not have a system. You have vibes. Would you replace a paid CRM with an open-source one? Write your take in the comments. More practical open-source solutions are on AtlasRepo in the channel header."),
    ("Postiz", "Did you know you can stop manually posting the same content everywhere? I recently found Postiz, an open-source repo for planning, scheduling, and distributing content across platforms. Calendar, composer, queue, connected channels — basically the boring part that decides whether anyone sees what you made. Would you use an open-source social media scheduler instead of a paid one? Tell me in the comments. And if you want more tools like this, AtlasRepo is linked in the channel header."),
]


audit_notes = """
Audit pass one: retention and clarity. The old weak point was that some ideas started too technically, as if the viewer already cared about the repo. This version starts with normal human context first: repetitive browser work, no editor, rented AI APIs, terminal friction, scattered leads, manual distribution. Then it introduces the repo as the answer. That makes the hooks easier to understand in the first three seconds.

Audit pass two: producer view. The order is now built for attention, not just technical logic. Video 1 starts with visual wow, then foundation, then business stack, then distribution. Video 2 turns AtlasRepo into an action layer, not just a directory. Video 3 uses a repeatable “job to be done” format so AtlasRepo can become a recurring series. The short scripts now use different closing moves: worker reveal, uncomfortable question, ownership frame, bridge frame, speed frame, rabbit-hole frame, review-the-diff frame, watchlist frame, grown-up business frame, and distribution frame.

Remaining production risk: Multica, MoneyPrinterTurbo, and OpenMontage need strong demo footage. If a live demo is hard, the partner should capture GitHub, docs, examples, product screenshots, and one clear proof screen. Do not force fake demos. If a repo cannot produce visual proof, keep it as “test” and use the limitation as trust-building narration.
"""


partner_tasks = {
    "General capture rules": [
        "Record every repo in this sequence: GitHub header, README promise, install/demo proof, product interface or output, limitation/risk, final visual verdict.",
        "For each repo, capture one horizontal version for long-form and one clean vertical crop for Shorts.",
        "Avoid private tokens, paid account data, API keys, customer names, browser history, and local file paths with sensitive information.",
        "The best shot is not a static README. The best shot is problem to action to result.",
    ],
    "browser-use": [
        "Goal: prove that AI can use websites, not just answer in a chat.",
        "Capture: GitHub header at https://github.com/browser-use/browser-use; README/demo; agent task; browser opening; clicking/typing/navigating; final result.",
        "Short focus: task prompt, visible browser actions, result.",
    ],
    "MoneyPrinterTurbo": [
        "Goal: show that AI short-video drafts can come from an open-source repo.",
        "Capture: GitHub header at https://github.com/harry0703/MoneyPrinterTurbo; input/topic screen; generation flow; final video/output; quality limitation if visible.",
        "Short focus: idea goes in, video draft comes out.",
    ],
    "Ollama": [
        "Goal: show local AI running on the machine.",
        "Capture: GitHub header at https://github.com/ollama/ollama; install/run command; model list or pull; one local prompt and answer.",
        "Short focus: local model answering outside a cloud chatbot UI.",
    ],
    "Open WebUI": [
        "Goal: show that local/cloud models can have a normal ChatGPT-like workspace.",
        "Capture: GitHub header at https://github.com/open-webui/open-webui; main UI; model selector; chat; settings or connection screen.",
        "Short focus: Ollama engine plus Open WebUI cockpit.",
    ],
    "Dify": [
        "Goal: show AI app building without starting from a blank folder.",
        "Capture: GitHub header at https://github.com/langgenius/dify; app builder; workflow or assistant setup; knowledge/RAG screen; published app result.",
        "Short focus: idea to usable AI app screen.",
    ],
    "ComfyUI": [
        "Goal: show visual AI workflow production.",
        "Capture: GitHub header at https://github.com/Comfy-Org/ComfyUI; node graph; prompt/model nodes; generated output; example gallery if useful.",
        "Short focus: node workflow to final visual.",
    ],
    "OpenHands": [
        "Goal: show an AI software engineer working inside a repo.",
        "Capture: GitHub header at https://github.com/OpenHands/OpenHands; task screen; file inspection; diff/change; result or test output.",
        "Short focus: task, code change, result.",
    ],
    "Multica": [
        "Goal: show AI coding agents managed like teammates.",
        "Capture: GitHub header at https://github.com/multica-ai/multica; dashboard; agents; task assignment; progress/status; any completed output.",
        "Short focus: one person managing several AI workers.",
    ],
    "Twenty CRM": [
        "Goal: show the business layer of a one-person company.",
        "Capture: GitHub header at https://github.com/twentyhq/twenty; CRM dashboard; contacts; companies; pipeline/deals; settings if clean.",
        "Short focus: scattered leads become a real pipeline.",
    ],
    "Postiz": [
        "Goal: show content distribution after one video is created. Context line for the edit: imagine one open-source repo, with a free/self-hostable path, can help distribute a single recorded video across multiple platforms.",
        "Capture: GitHub header at https://github.com/gitroomhq/postiz-app; calendar; composer; queue; connected channels; scheduled post.",
        "Short focus: one video becomes scheduled content across platforms.",
    ],
    "Video 2 micro-SaaS demo": [
        "Goal: create or mock a simple Repo Radar screen if the real build is not ready.",
        "Capture: niche input; repo cards; save/test/skip verdicts; generated Shorts ideas; generated partner brief.",
        "Required example card: Postiz with the context-first hook about one video being distributed across platforms.",
    ],
    "Video 3 AtlasRepo demo": [
        "Goal: make AtlasRepo feel like a discovery method, not a random directory.",
        "Capture: atlasrepo.com homepage/search; query or category for content factory; cards for Postiz, MoneyPrinterTurbo, OpenMontage; click-through to GitHub.",
        "If OpenMontage demo is hard: capture README, docs, examples, and explain it as a deeper test rather than a guaranteed ready-to-use product.",
    ],
}


def make_scripts_doc():
    doc = setup_doc(
        "AtlasRepo First Batch Scripts — Context-First US English",
        "Read-aloud long-form scripts and Shorts scripts. Context-first hooks, slippery-slope pacing, native AtlasRepo mentions.",
    )
    doc.add_heading("Producer Note", level=1)
    doc.add_paragraph(
        "These scripts are written to be read out loud. The pattern is: normal-life context first, repo second, proof third, verdict last. Do not read section labels on camera.",
        style="ScriptBody",
    )
    add_script(doc, "Long Video 1 — These 10 GitHub Repos Can Build a One-Person AI Company", "Target length: 10-15 minutes with demos and pauses.", video1)
    add_script(doc, "Long Video 2 — I Built a Micro-SaaS in One Evening with Open-Source Projects", "Target length: 10-15 minutes with product/mockup demo.", video2)
    add_script(doc, "Long Video 3 — This Website Found 3 Gold GitHub Repos for a Content Factory", "Target length: 10-12 minutes with AtlasRepo walkthrough.", video3)
    doc.add_heading("10 Shorts — Read-Aloud Scripts", level=1)
    for idx, (repo, script) in enumerate(shorts, 1):
        add_script(doc, f"Short {idx} — {repo}", "Target length: 35-55 seconds.", script)
    add_script(doc, "Self-Audit Notes", "Retention and producer audit after regeneration.", audit_notes)
    return doc


def make_partner_doc():
    doc = setup_doc(
        "AtlasRepo First Batch — Partner Screencast Brief",
        "Exact capture requests for 3 long-form videos and 10 Shorts. US English version.",
    )
    doc.add_heading("Recording Principle", level=1)
    doc.add_paragraph(
        "Every clip should prove a simple context-first claim. Show the human problem, then the GitHub repo, then the screen proof that the repo can help with that problem.",
        style="BriefBody",
    )
    for title, lines in partner_tasks.items():
        add_brief_section(doc, title, lines)
    return doc


if __name__ == "__main__":
    scripts_path = OUT / "atlasrepo_first_batch_read_aloud_scripts_EN_context.docx"
    partner_path = OUT / "atlasrepo_first_batch_partner_screencast_brief_EN_context.docx"
    make_scripts_doc().save(scripts_path)
    make_partner_doc().save(partner_path)
    print(scripts_path)
    print(partner_path)
